"""Resume Tailor: suggest new bullets to ADD (never rewrite existing text) against a pasted
JD, with a live ATS keyword score and a downloadable tailored .docx. Ported directly from
studysage-rag/src/resume_tailor.py - identical logic, only the LLM call now goes through
app.services.llm_provider.get_llm(provider=...) instead of a session-state-driven default."""
import re
import difflib
import io
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from docx import Document as DocxDocument
from docx.shared import Pt
from app.services.llm_provider import get_llm, invoke_and_check_truncation


_ATS_STOPWORDS = {
    "the", "a", "an", "and", "or", "to", "for", "of", "in", "on", "with", "using", "at", "by",
    "from", "as", "was", "were", "is", "are", "that", "this", "into", "be", "will", "you", "your",
    "we", "our", "us", "job", "role", "work", "working", "team", "years", "year", "experience",
    "strong", "ability", "skills", "skill", "knowledge", "including", "etc", "must", "should",
    "responsibilities", "requirements", "required", "preferred", "plus", "about", "who", "what",
    "have", "has", "had", "can", "all", "any", "other", "such", "not", "than", "their", "they",
    "it", "its", "these", "those", "also", "new", "across", "within", "per", "based", "related",
    "looking", "candidate", "candidates", "company", "position", "opportunity", "description",
    "hiring", "infrastructure", "expertise", "code", "environment", "environments", "systems",
    "system", "solutions", "solution", "large", "high", "level", "levels", "join", "great",
}


def extract_jd_keywords(jd_text, max_keywords=30):
    """Cheap, deterministic keyword extraction (no LLM call, so it can run on every checkbox
    click for a truly 'live' score) - pulls out the tool/tech/skill terms an ATS keyword scan
    would actually look for: capitalized/CamelCase tokens (tool names), acronyms, and
    multi-word technical phrases, ranked by frequency."""
    phrase_pattern = re.compile(r"\b([A-Z][a-zA-Z0-9+#]*(?:\s+[A-Z][a-zA-Z0-9+#]*){0,2})\b")
    candidates = phrase_pattern.findall(jd_text)

    freq = {}
    for c in candidates:
        key = c.strip().rstrip(".")
        if len(key) < 2:
            continue
        low = key.lower()
        if low in _ATS_STOPWORDS:
            continue
        if all(w.lower() in _ATS_STOPWORDS for w in key.split()):
            continue
        freq[key] = freq.get(key, 0) + 1

    for w in re.findall(r"\b[a-z][a-z0-9+#./-]{2,}\b", jd_text.lower()):
        if w in _ATS_STOPWORDS or len(w) < 3:
            continue
        freq.setdefault(w, 0)
        freq[w] += 0.3

    ranked = sorted(freq.items(), key=lambda kv: (-kv[1], kv[0]))
    seen_lower = set()
    keywords = []
    for term, _ in ranked:
        low = term.lower()
        if low in seen_lower:
            continue
        seen_lower.add(low)
        keywords.append(term)
        if len(keywords) >= max_keywords:
            break
    return keywords


def ats_match_score(resume_text, jd_keywords):
    """Returns (score 0-100, matched keywords, missing keywords)."""
    if not jd_keywords:
        return 0, [], []
    resume_low = resume_text.lower()
    matched, missing = [], []
    for kw in jd_keywords:
        if kw.lower() in resume_low:
            matched.append(kw)
        else:
            missing.append(kw)
    score = round(100 * len(matched) / len(jd_keywords))
    return score, matched, missing


def analyze_resume_for_jd(resume_text, jd_text, provider: str = "groq"):
    """Finds the first two project sections in the resume and, against the JD, produces
    SUGGESTIONS (new bullets to ADD, existing text never touched) and MISSING_SKILLS."""
    llm = get_llm(provider=provider, temperature=0.4, max_tokens=1800)
    template = """You are a resume coach. Below is a candidate's RESUME and a JOB DESCRIPTION.

RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text}

Find the FIRST TWO distinct project/experience entries in the resume (in the order they
appear). For EACH of the two, produce:

1. ORIGINAL: the verbatim text of that project section, copied exactly from the resume, so it
   can be matched back against the original document. Do NOT reword or restructure this — it
   must be an exact copy.

2. SUGGESTIONS: 3-5 NEW candidate bullet points that could be ADDED to this project (never
   modifications to existing lines — existing text is never touched). Each must be a ready-to-
   insert, first-person, action-verb resume bullet, plausible from context and relevant to the
   JD, never claiming a tool/platform the candidate has no evidence of using. Keep each bullet
   to one concise sentence - no filler.

   For EACH suggestion, decide individually whether a quantitative metric would genuinely
   strengthen THAT specific point (e.g. a bullet about optimization or scale usually benefits
   from a number; a bullet about collaboration or process usually doesn't). If yes, weave a
   plausible ILLUSTRATIVE ESTIMATED RANGE directly into that bullet's sentence (e.g. "...cutting
   processing time by roughly 25-30%"). If a metric doesn't naturally fit that particular point,
   leave it out entirely — do not force a number into every bullet. These ranges are estimates
   for the candidate to review and replace with their real figure, never confirmed facts.

3. MISSING_SKILLS: for each tool/language/platform the JOB DESCRIPTION asks for that is NOT
   evidenced anywhere in this project or the resume, output ONE LINE in this exact pipe format:
   SKILL :: NOTE :: DRAFT
   - SKILL: the tool/tech name
   - NOTE: one sentence flagging this as a real gap, not something to fabricate
   - DRAFT: a draft bullet blending this skill into the REAL project context ONLY IF plausible
     (phrase speculatively, e.g. "Could describe: ..."). If no plausible tie-in exists, DRAFT
     should say "No plausible tie-in — do not add this."
   If no missing skills for this project, output: None :: None :: None

Respond in EXACTLY this format, nothing else:

PROJECT1_ORIGINAL:
<verbatim text>
PROJECT1_SUGGESTIONS:
- <suggestion, with an embedded metric range only if it genuinely fits>
PROJECT1_MISSING_SKILLS:
<SKILL :: NOTE :: DRAFT line(s), or "None :: None :: None">
===
PROJECT2_ORIGINAL:
<verbatim text>
PROJECT2_SUGGESTIONS:
- <suggestion, with an embedded metric range only if it genuinely fits>
PROJECT2_MISSING_SKILLS:
<SKILL :: NOTE :: DRAFT line(s), or "None :: None :: None">

Begin now:"""
    prompt = ChatPromptTemplate.from_template(template)
    messages = prompt.format_messages(resume_text=resume_text, jd_text=jd_text)
    result, truncated = invoke_and_check_truncation(llm, messages)
    projects = _parse_projects(result)
    return projects, truncated


_METRIC_PATTERN = re.compile(r"\d+\s*[-–]\s*\d+\s*%|\d+%|\d[\d,]*\s*(records|requests|users|gb|tb|ms|/day|/sec)", re.IGNORECASE)


def _bullets(text):
    return [line.strip("- ").strip() for line in text.strip().splitlines() if line.strip().startswith("-")]


def _parse_missing_skills(text):
    skills = []
    for line in text.strip().splitlines():
        line = line.strip().lstrip("-").strip()
        if not line or "::" not in line:
            continue
        parts = [p.strip() for p in line.split("::")]
        if len(parts) != 3:
            continue
        skill, note, draft = parts
        if skill.lower() == "none":
            continue
        skills.append({"skill": skill, "note": note, "draft": draft})
    return skills


def _parse_projects(raw):
    blocks = raw.split("===")
    projects = []
    pattern = re.compile(
        r"PROJECT\d+_ORIGINAL:\s*(?P<original>.*?)\s*"
        r"PROJECT\d+_SUGGESTIONS:\s*(?P<suggestions>.*?)\s*"
        r"PROJECT\d+_MISSING_SKILLS:\s*(?P<missing>.*)",
        re.DOTALL
    )
    for block in blocks:
        match = pattern.search(block)
        if not match:
            continue
        suggestions = _bullets(match.group("suggestions"))
        projects.append({
            "original": match.group("original").strip(),
            "suggestions": suggestions,
            "has_any_metric": any(_METRIC_PATTERN.search(s) for s in suggestions),
            "missing_skills": _parse_missing_skills(match.group("missing")),
        })
    return projects


def _word_diff_line(orig_line, edit_line):
    orig_words, edit_words = orig_line.split(), edit_line.split()
    sm = difflib.SequenceMatcher(None, orig_words, edit_words)
    left_parts, right_parts = [], []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            left_parts.append(" ".join(orig_words[i1:i2]))
            right_parts.append(" ".join(edit_words[j1:j2]))
        elif tag == "delete":
            left_parts.append(
                f'<span style="background:#4a1414;color:#ff8a8a;text-decoration:line-through;padding:1px 3px;border-radius:3px;">{" ".join(orig_words[i1:i2])}</span>'
            )
        elif tag == "insert":
            right_parts.append(
                f'<span style="background:#143d1d;color:#8affa0;padding:1px 3px;border-radius:3px;">{" ".join(edit_words[j1:j2])}</span>'
            )
        elif tag == "replace":
            left_parts.append(
                f'<span style="background:#4a1414;color:#ff8a8a;text-decoration:line-through;padding:1px 3px;border-radius:3px;">{" ".join(orig_words[i1:i2])}</span>'
            )
            right_parts.append(
                f'<span style="background:#143d1d;color:#8affa0;padding:1px 3px;border-radius:3px;">{" ".join(edit_words[j1:j2])}</span>'
            )
    return " ".join(left_parts), " ".join(right_parts)


def render_side_by_side_diff(original, edited):
    """GitHub PR-style split diff: returns (left_html, right_html). Diffs LINE by line so
    bullet-point structure is preserved."""
    orig_lines = list(original.splitlines())
    edit_lines = list(edited.splitlines())
    sm = difflib.SequenceMatcher(None, orig_lines, edit_lines)

    left_rows, right_rows = [], []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for line in orig_lines[i1:i2]:
                left_rows.append(line)
                right_rows.append(line)
        elif tag == "delete":
            for line in orig_lines[i1:i2]:
                left_rows.append(f'<span style="background:#4a1414;color:#ff8a8a;text-decoration:line-through;padding:1px 3px;border-radius:3px;">{line}</span>')
        elif tag == "insert":
            for line in edit_lines[j1:j2]:
                right_rows.append(f'<span style="background:#143d1d;color:#8affa0;padding:1px 3px;border-radius:3px;">{line}</span>')
        elif tag == "replace":
            o_block, e_block = orig_lines[i1:i2], edit_lines[j1:j2]
            paired = min(len(o_block), len(e_block))
            for k in range(paired):
                l, r = _word_diff_line(o_block[k], e_block[k])
                left_rows.append(l)
                right_rows.append(r)
            for line in o_block[paired:]:
                left_rows.append(f'<span style="background:#4a1414;color:#ff8a8a;text-decoration:line-through;padding:1px 3px;border-radius:3px;">{line}</span>')
            for line in e_block[paired:]:
                right_rows.append(f'<span style="background:#143d1d;color:#8affa0;padding:1px 3px;border-radius:3px;">{line}</span>')

    def _as_html(rows):
        out = []
        for row in rows:
            stripped = row.strip()
            if not stripped:
                out.append("<div style='height:6px'></div>")
            else:
                out.append(f"<div style='margin:4px 0;'>&bull;&nbsp; {row}</div>")
        return "\n".join(out)

    return _as_html(left_rows), _as_html(right_rows)


def insert_naturally(original, new_bullets):
    """Insert each approved new bullet right after the existing line it's most topically
    related to (by word overlap), instead of always tacking everything onto the end."""
    if not new_bullets:
        return original

    lines = original.splitlines()
    stopwords = {
        "the", "a", "an", "and", "or", "to", "for", "of", "in", "on", "with", "using",
        "at", "by", "from", "as", "was", "were", "is", "are", "that", "this", "into"
    }

    def _keywords(text):
        return {w.strip(".,;:()").lower() for w in text.split() if len(w) > 3 and w.lower() not in stopwords}

    line_keywords = [_keywords(l) for l in lines]
    insertions = {}
    tail = []

    for bullet in new_bullets:
        b_keywords = _keywords(bullet)
        best_idx, best_score = None, 0
        for idx, lk in enumerate(line_keywords):
            if not lk:
                continue
            overlap = len(b_keywords & lk)
            if overlap > best_score:
                best_score, best_idx = overlap, idx
        if best_idx is not None and best_score > 0:
            insertions.setdefault(best_idx, []).append(bullet)
        else:
            tail.append(bullet)

    result = []
    for idx, line in enumerate(lines):
        result.append(line)
        for bullet in insertions.get(idx, []):
            result.append(bullet)
    result.extend(tail)
    return "\n".join(result)


def _flexible_find(snippet, haystack):
    """Whitespace-tolerant search since PDF/DOCX text extraction often reflows whitespace."""
    snippet_lines = [l.strip() for l in snippet.splitlines() if l.strip()]
    if not snippet_lines:
        return None
    pattern_parts = [re.escape(l) for l in snippet_lines]
    pattern = r"\s+".join(pattern_parts)
    try:
        match = re.search(pattern, haystack, flags=re.DOTALL)
    except re.error:
        return None
    return match


_SUBHEADER_PATTERN = re.compile(r"^(Client|Role|Company|Project|Title)\s*:", re.IGNORECASE)


def _classify_lines(text):
    """Mirrors typical resume structure: name + contact as the first two non-blank lines,
    ALL-CAPS short lines as section headers, 'Client:'/'Role:' lines as bold sub-headers,
    tab-indented or long lines as bullets. Yields (kind, clean_text) per non-blank line."""
    seen_nonblank = 0
    for raw_line in text.split("\n"):
        stripped_raw = raw_line.strip()
        clean = stripped_raw.lstrip("-•*").strip()
        if not clean:
            continue
        seen_nonblank += 1
        if seen_nonblank == 1:
            yield "name", clean
        elif seen_nonblank == 2 and ("|" in clean or "@" in clean or re.search(r"\d{3}", clean)):
            yield "contact", clean
        elif _SUBHEADER_PATTERN.match(clean):
            yield "subheader", clean
        elif raw_line.startswith("\t\t"):
            yield "subbullet", clean
        elif raw_line.startswith("\t") or len(clean) > 45:
            yield "bullet", clean
        elif clean.isupper() and len(clean) < 45:
            yield "header", clean
        else:
            yield "subheader", clean


def _write_body(doc, text):
    for kind, clean in _classify_lines(text):
        if kind == "name":
            p = doc.add_paragraph()
            p.alignment = 1
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(16)
        elif kind == "contact":
            p = doc.add_paragraph()
            p.alignment = 1
            run = p.add_run(clean)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(8)
        elif kind == "header":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(12)
        elif kind == "subheader":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(clean)
            run.bold = True
        elif kind == "subbullet":
            p = doc.add_paragraph(clean, style="List Bullet 2")
            p.paragraph_format.space_after = Pt(0)
        else:
            p = doc.add_paragraph(clean, style="List Bullet")
            p.paragraph_format.space_after = Pt(0)


def build_tailored_docx(full_resume_text, replacements):
    """replacements: list of (original_snippet, final_text) tuples. Returns a BytesIO buffer."""
    final_text = full_resume_text
    unmatched = []
    for original_snippet, final_snippet in replacements:
        if original_snippet in final_text:
            final_text = final_text.replace(original_snippet, final_snippet, 1)
            continue
        match = _flexible_find(original_snippet, final_text)
        if match:
            final_text = final_text[:match.start()] + final_snippet + final_text[match.end():]
        else:
            unmatched.append(final_snippet)

    doc = DocxDocument()
    _write_body(doc, final_text)

    if unmatched:
        doc.add_paragraph("")
        doc.add_paragraph("--- Tailored additions (could not auto-place in original text) ---")
        for snippet in unmatched:
            _write_body(doc, snippet)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def render_tailored_preview_markdown(full_resume_text, replacements):
    final_text = full_resume_text
    unmatched = []
    for original_snippet, final_snippet in replacements:
        if original_snippet in final_text:
            final_text = final_text.replace(original_snippet, final_snippet, 1)
            continue
        match = _flexible_find(original_snippet, final_text)
        if match:
            final_text = final_text[:match.start()] + final_snippet + final_text[match.end():]
        else:
            unmatched.append(final_snippet)

    lines_md = []
    for kind, clean in _classify_lines(final_text):
        if kind == "name":
            lines_md.append(f"### {clean}")
        elif kind == "contact":
            lines_md.append(f"*{clean}*")
        elif kind == "header":
            lines_md.append(f"\n**{clean.upper()}**")
        elif kind == "subheader":
            lines_md.append(f"**{clean}**")
        elif kind == "subbullet":
            lines_md.append(f"    - {clean}")
        else:
            lines_md.append(f"- {clean}")

    if unmatched:
        lines_md.append("\n---\n*Tailored additions (could not auto-place):*")
        for snippet in unmatched:
            for line in snippet.split("\n"):
                clean = line.strip()
                if clean:
                    lines_md.append(f"- {clean}")

    return "\n".join(lines_md)
